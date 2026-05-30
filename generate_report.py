"""
Sinh file báo cáo Word cho đồ án EZWare_Backend.

Bám sát Template_IE221_New.docx:
  H1: GIỚI THIỆU (không đổi tên)
  H1: MÔ TẢ CƠ SỞ DỮ LIỆU
  H1: TIẾP THEO NỮA — Chức năng hệ thống
  H1: THIẾT KẾ HỆ THỐNG
  H1: KIẾN TRÚC TRIỂN KHAI
  H1: KẾT LUẬN
  H1: TÀI LIỆU THAM KHẢO (style TLTK — sai trừ 1 điểm)
  PHỤ LỤC A. PHÂN CÔNG NHIỆM VỤ
  PHỤ LỤC B. TÓM TẮT KẾT QUẢ        (BẮT BUỘC)
  PHỤ LỤC C. DANH SÁCH ENDPOINT API
  PHỤ LỤC D. CHI TIẾT SCHEMA CSDL
  PHỤ LỤC E. HƯỚNG DẪN CÀI ĐẶT VÀ CHẠY
  PHỤ LỤC F. HÌNH KIỂM THỬ
  PHỤ LỤC G. CODE TRÍCH (tuỳ chọn)

Footer (chỉ section 2 trở đi, bỏ qua trang bìa):
  • Trái: tên SV viết tắt (placeholder Nguyễn V.A — Nguyễn V.B)
  • Phải: Trang X

Chạy:
    python generate_report.py
Output: NhomXX_Bao_cao.docx

Yêu cầu: pip install python-docx
"""
from datetime import datetime

import os

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Thư mục chứa các PNG sinh từ generate_diagrams.py
IMG_DIR = 'report_images'


OUTPUT_NAME = 'NhomXX_Bao_cao_v4.docx'

# URL public deploy thực tế trên Render — sử dụng nhất quán trong báo cáo
DEPLOY_URL = 'https://ezware-backend.onrender.com'

# Placeholder thông tin nhóm — sinh viên thay sau khi mở file
GROUP_NUM = 'Nhóm XX'
MEMBERS = [
    # (STT, Họ tên, MSSV, Ngành)
    ('1', 'Nguyễn Văn A', '25XXXXX1', 'CNTT'),
    ('2', 'Nguyễn Văn B', '25XXXXX2', 'CNTT'),
]
# Tên rút gọn cho footer (template gợi ý: Nguyễn Thị Thu Trang => Nguyễn T.T. Trang)
MEMBERS_SHORT = 'Nguyễn V.A — Nguyễn V.B'


# ============================================================
# Helpers
# ============================================================
def set_cell_border(cell, color='000000', size='4'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_section_page_break(doc):
    """Section break thay vì page break thường — để footer mới có hiệu lực."""
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.top_margin = Cm(2.0)
    new_section.bottom_margin = Cm(2.0)
    new_section.left_margin = Cm(2.5)
    new_section.right_margin = Cm(2.0)
    return new_section


def add_centered(doc, text, bold=False, size=13, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_bt(doc, text, first_line_indent=1.0):
    """Đoạn body text style BT."""
    p = doc.add_paragraph(style='BT')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    return p


def add_g1(doc, text):
    """Bullet style G1."""
    p = doc.add_paragraph(style='G1')
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f'•  {text}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    return p


def add_heading_1(doc, text):
    p = doc.add_paragraph(style='Heading 1')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    return p


def add_heading_2(doc, text):
    p = doc.add_paragraph(style='Heading 2')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    return p


def add_tltk(doc, text):
    p = doc.add_paragraph(style='TLTK')
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


def add_figure(doc, image_file, caption, width_inches=6.3):
    """Chèn hình PNG + caption tiếng Việt. Fallback về placeholder text nếu file không tồn tại."""
    img_path = os.path.join(IMG_DIR, image_file)
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Inches(width_inches))
    else:
        # Fallback nếu chưa chạy generate_diagrams.py
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'[CHÈN HÌNH TẠI ĐÂY — {caption}]')
        run.italic = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(120, 120, 120)
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    cap_run.italic = True
    cap_run.font.name = 'Times New Roman'
    cap_run.font.size = Pt(12)


# Backwards-compatible alias (cũ gọi placeholder, giờ thử insert hình)
def add_figure_placeholder(doc, caption):
    # Tự suy ra tên file từ caption: "Hình 1." -> hinh1_*
    if caption.lower().startswith('hình 1'):
        add_figure(doc, 'hinh1_er_diagram.png', caption)
    elif caption.lower().startswith('hình 2'):
        add_figure(doc, 'hinh2_function_diagram.png', caption)
    elif caption.lower().startswith('hình 3'):
        add_figure(doc, 'hinh3_class_diagram.png', caption)
    elif caption.lower().startswith('hình 4'):
        add_figure(doc, 'hinh4_postman_login.png', caption)
    elif caption.lower().startswith('hình 5'):
        add_figure(doc, 'hinh5_postman_receipt_flow.png', caption)
    elif caption.lower().startswith('hình 6'):
        add_figure(doc, 'hinh6_swagger_ui.png', caption)
    else:
        add_figure(doc, '__not_exist__.png', caption)


def fix_table_font(t, size=11, bold_header=True):
    """Đặt font Times New Roman cho toàn bộ table + border."""
    for i, row in enumerate(t.rows):
        for c in row.cells:
            set_cell_border(c)
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(size)
                    if bold_header and i == 0:
                        run.bold = True


# ============================================================
# Styles
# ============================================================
def create_custom_styles(doc):
    styles = doc.styles
    existing = [s.name for s in styles]
    if 'BT' not in existing:
        bt = styles.add_style('BT', 1)
        bt.base_style = styles['Normal']
        bt.font.name = 'Times New Roman'
        bt.font.size = Pt(13)
        bt.paragraph_format.line_spacing = 1.3
        bt.paragraph_format.space_after = Pt(4)
    if 'G1' not in existing:
        g1 = styles.add_style('G1', 1)
        g1.base_style = styles['Normal']
        g1.font.name = 'Times New Roman'
        g1.font.size = Pt(13)
    if 'TLTK' not in existing:
        tltk = styles.add_style('TLTK', 1)
        tltk.base_style = styles['Normal']
        tltk.font.name = 'Times New Roman'
        tltk.font.size = Pt(12)


# ============================================================
# Footer with page number
# ============================================================
def add_page_number_field(paragraph):
    """Chèn field PAGE để Word tự render số trang."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def setup_footer(section, members_short):
    """Footer 2 cột: trái = tên SV, phải = Trang X."""
    footer = section.footer
    # Reset paragraph hiện có
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Dùng tab stop để chia trái/phải
    p.paragraph_format.tab_stops.add_tab_stop(
        Cm(16.0), alignment=WD_ALIGN_PARAGRAPH.RIGHT
    )
    # Trái: tên SV
    run_left = p.add_run(members_short)
    run_left.font.name = 'Times New Roman'
    run_left.font.size = Pt(10)
    # Tab sang phải
    run_tab = p.add_run('\t')
    run_tab.font.name = 'Times New Roman'
    # Phải: Trang X
    run_right = p.add_run('Trang ')
    run_right.font.name = 'Times New Roman'
    run_right.font.size = Pt(10)
    add_page_number_field(p)
    # Resize tất cả run trong p về 10pt + Times
    for run in p.runs:
        run.font.name = 'Times New Roman'
        if not run.font.size:
            run.font.size = Pt(10)


# ============================================================
# Section content builders
# ============================================================
def build_cover(doc):
    """Trang bìa."""
    add_centered(doc, 'ĐẠI HỌC QUỐC GIA TP. HỒ CHÍ MINH', bold=True, size=14)
    add_centered(doc, 'TRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN', bold=True, size=14)
    for _ in range(2):
        doc.add_paragraph()

    add_centered(doc, 'ĐỒ ÁN MÔN HỌC IE221', bold=True, size=18)
    add_centered(doc, '(LẬP TRÌNH PYTHON)', italic=True, size=14)
    for _ in range(2):
        doc.add_paragraph()

    add_centered(doc, 'HỆ THỐNG QUẢN LÝ KHO EZWARE', bold=True, size=20)
    add_centered(doc, 'XÂY DỰNG BACKEND API VỚI', bold=True, size=20)
    add_centered(doc, 'DJANGO REST FRAMEWORK', bold=True, size=20)
    for _ in range(2):
        doc.add_paragraph()

    # Bảng nhóm — match Table 0 của template
    # Row 0: Nhóm XX (merge 4 cells)
    # Row 1: Sinh viên thực hiện... (merge 4 cells)
    # Row 2: STT | Họ tên | MSSV | Ngành (header)
    # Row 3-7: 5 student rows
    total_rows = 3 + 5  # header + 5 students slots
    t = doc.add_table(rows=total_rows, cols=4)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Hàng 0: Nhóm XX — merge 4 cells
    r0 = t.rows[0].cells
    r0[0].merge(r0[1]).merge(r0[2]).merge(r0[3])
    r0[0].text = GROUP_NUM
    r0[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0[0].paragraphs[0].runs[0].bold = True
    r0[0].paragraphs[0].runs[0].font.size = Pt(14)

    # Hàng 1: Sinh viên thực hiện — merge 4 cells
    r1 = t.rows[1].cells
    r1[0].merge(r1[1]).merge(r1[2]).merge(r1[3])
    r1[0].text = 'Sinh viên thực hiện:'
    r1[0].paragraphs[0].runs[0].bold = True

    # Hàng 2: header
    headers = ['STT', 'Họ tên', 'MSSV', 'Ngành']
    for i, h in enumerate(headers):
        t.rows[2].cells[i].text = h
        t.rows[2].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        t.rows[2].cells[i].paragraphs[0].runs[0].bold = True

    # Hàng 3+: SV (đổ data MEMBERS rồi để trống slot dư)
    for i in range(5):
        row_cells = t.rows[3 + i].cells
        if i < len(MEMBERS):
            sv = MEMBERS[i]
        else:
            sv = (str(i + 1), '', '', '')
        for j, val in enumerate(sv):
            row_cells[j].text = val
            row_cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    fix_table_font(t, size=12, bold_header=False)
    # Header hàng 0, 1, 2 in đậm — đã set thủ công bên trên
    for c in t.rows[0].cells + t.rows[1].cells + t.rows[2].cells:
        for p in c.paragraphs:
            for run in p.runs:
                run.bold = True

    for _ in range(3):
        doc.add_paragraph()

    add_centered(
        doc, f'TP. HỒ CHÍ MINH – {datetime.now().strftime("%m/%Y")}',
        bold=True, size=13,
    )


def build_intro(doc):
    add_heading_1(doc, '1. GIỚI THIỆU')
    add_bt(doc,
           'Đề tài “Hệ thống quản lý kho EZWare” là một ứng dụng backend cung '
           'cấp REST API phục vụ nghiệp vụ quản lý kho hàng cho doanh nghiệp '
           'nhỏ và vừa. Hệ thống cho phép quản trị viên và nhân viên kho thực '
           'hiện các thao tác: quản lý danh mục sản phẩm, quản lý danh mục kho, '
           'lập phiếu nhập / xuất kho, duyệt phiếu để cập nhật tồn kho thực tế, '
           'và tra cứu báo cáo tồn kho thấp. Chúng tôi sử dụng ngôn ngữ Python '
           'kết hợp với framework Django 5 và Django REST Framework để xây '
           'dựng các endpoint RESTful theo chuẩn HTTP (GET / POST / PUT / DELETE), '
           'kèm cơ chế xác thực bằng Token và tài liệu API tự sinh thông qua '
           'thư viện drf-yasg (Swagger UI). Kết quả đạt được là một backend API '
           'hoàn chỉnh với 14 URL endpoint và 23 thao tác HTTP, 6 bảng cơ sở '
           'dữ liệu (SQLite), phân quyền theo hai vai trò (admin / staff), '
           'nghiệp vụ duyệt phiếu được bảo vệ bằng transaction nguyên tử kết '
           'hợp khóa hàng (row-level lock), bộ kiểm thử bằng Postman với hơn '
           f'30 test case, và đã được triển khai công khai tại {DEPLOY_URL}.')
    add_bt(doc,
           'Trong đề tài này, chúng tôi tự phân tích và thiết kế cơ sở dữ liệu '
           'cũng như toàn bộ các API mà không tham khảo nguyên mẫu nào trên '
           'Internet. Một số thư viện và mẫu thiết kế phổ biến của Django REST '
           'được áp dụng theo tài liệu chính thức của framework [1] [2].')


def build_database(doc):
    add_heading_1(doc, '2. MÔ TẢ CƠ SỞ DỮ LIỆU EZWARE')
    add_bt(doc,
           'Cơ sở dữ liệu của hệ thống được thiết kế trên SQLite, gồm 6 bảng có '
           'quan hệ chặt chẽ với nhau để mô tả các thực thể của nghiệp vụ kho '
           '(nằm trong khoảng 5–10 bảng theo yêu cầu của đồ án). Sơ đồ quan hệ '
           'giữa các bảng được trình bày tại Hình 1.')
    add_figure_placeholder(doc, 'Hình 1. Sơ đồ ER cơ sở dữ liệu EZWare.')

    add_bt(doc, 'Sáu bảng trong cơ sở dữ liệu bao gồm:')
    add_g1(doc, 'users — quản lý tài khoản người dùng với hai vai trò (admin, staff).')
    add_g1(doc, 'products — danh mục sản phẩm. Không lưu đơn giá trong bảng để giữ tính trung tính của module quản lý kho; mọi đơn vị tiền tệ trong nghiệp vụ tính theo VNĐ.')
    add_g1(doc, 'warehouses — danh mục các kho hàng kèm địa chỉ.')
    add_g1(doc, 'inventory_receipts — đầu phiếu nhập / xuất, có trạng thái PENDING / APPROVED / CANCELLED.')
    add_g1(doc, 'receipt_details — chi tiết các dòng sản phẩm trong từng phiếu.')
    add_g1(doc, 'inventory — bảng tồn kho thực tế, được cập nhật sau khi phiếu được duyệt; ràng buộc UNIQUE(warehouse_id, product_id) để mỗi sản phẩm chỉ có một dòng tồn ở mỗi kho.')

    add_heading_2(doc, '2.1. Người dùng và phân quyền')
    add_bt(doc,
           'Bảng users kế thừa từ AbstractBaseUser của Django, có cột user_role '
           'với hai giá trị: admin và staff. Admin được quyền cao nhất: quản lý '
           'danh mục, duyệt phiếu, xóa dữ liệu. Staff chỉ được phép xem danh '
           'mục và tạo phiếu / thêm chi tiết phiếu. Mật khẩu được lưu trữ ở '
           'dạng băm PBKDF2 (mặc định của Django) thay vì lưu plain-text. '
           'Hệ thống đã đảm bảo an toàn ở tầng đăng ký: không cho phép client '
           'tự chọn user_role qua endpoint /api/auth/register để tránh leo '
           'thang quyền — vai trò admin chỉ được tạo bởi admin sẵn có.')

    add_heading_2(doc, '2.2. Quan hệ giữa các bảng nghiệp vụ')
    add_bt(doc,
           'Bảng inventory_receipts kết nối với warehouses qua warehouse_id '
           '(khóa ngoại với on_delete=PROTECT — không cho phép xóa kho khi đã '
           'phát sinh phiếu); receipt_details liên kết với inventory_receipts '
           'qua receipt_id (CASCADE — xóa phiếu sẽ xóa toàn bộ chi tiết kèm '
           'theo). Bảng inventory không lấy trực tiếp giá trị từ phiếu mà được '
           'tính lại sau mỗi lần phiếu được duyệt, đảm bảo dữ liệu tồn kho '
           'luôn phản ánh chính xác các phiếu đã có hiệu lực.')


def build_features(doc):
    """H1 thứ 3 — TIẾP THEO NỮA, mô tả các chức năng và kiểm thử."""
    add_heading_1(doc, '3. CHỨC NĂNG HỆ THỐNG')
    add_bt(doc,
           'Hệ thống được tổ chức thành 4 nhóm chức năng chính tương ứng với 4 '
           'app Django. Sơ đồ chức năng tổng quan được trình bày tại Hình 2.')
    add_figure_placeholder(doc, 'Hình 2. Sơ đồ chức năng tổng quan của hệ thống EZWare.')

    add_g1(doc, 'Xác thực và quản lý tài khoản: đăng ký, đăng nhập, đăng xuất, xem / cập nhật thông tin cá nhân (4 endpoint, app accounts).')
    add_g1(doc, 'Quản lý danh mục: CRUD sản phẩm và CRUD kho hàng; phân quyền theo IsAdminOrReadOnly — staff đọc được, admin sửa được (5 endpoint, app products và warehouses).')
    add_g1(doc, 'Vận hành kho: tạo phiếu, thêm chi tiết, duyệt / hủy phiếu (4 endpoint, app inventory). Đây là nhóm có business logic phức tạp nhất.')
    add_g1(doc, 'Báo cáo: xem tồn kho theo kho và cảnh báo các sản phẩm tồn thấp theo ngưỡng (2 endpoint, app reports).')

    add_heading_2(doc, '3.1. Cơ chế xác thực và phân quyền')
    add_bt(doc,
           'Hệ thống sử dụng Token Authentication của DRF. Sau khi đăng nhập '
           'thành công, client nhận một token và gửi kèm header '
           '“Authorization: Token <key>” ở mọi request tiếp theo. Phân quyền '
           'được thực thi bằng hai lớp Permission tự định nghĩa kế thừa '
           'BasePermission: IsAdminRole (chỉ admin được phép) và IsAdminOrReadOnly '
           '(mọi user đăng nhập đều đọc được, nhưng chỉ admin được phép '
           'thay đổi dữ liệu master).')

    add_heading_2(doc, '3.2. Bảo vệ tính nhất quán khi duyệt phiếu')
    add_bt(doc,
           'Đây là chức năng quan trọng và được xem là điểm hài lòng nhất của '
           'hệ thống. Khi admin duyệt một phiếu xuất, hệ thống lần lượt thực '
           'hiện: (1) kiểm tra tồn kho hiện tại của từng sản phẩm trong phiếu '
           'có đủ không — nếu không đủ thì trả 400 với thông báo chi tiết; '
           '(2) nếu đủ, mở một transaction nguyên tử (transaction.atomic), '
           'trong đó sử dụng select_for_update để khóa các dòng tồn kho liên '
           'quan, sau đó trừ số lượng tương ứng; (3) cập nhật trạng thái phiếu '
           'sang APPROVED. Cơ chế này đảm bảo: nếu hai admin cùng duyệt hai '
           'phiếu xuất cùng một sản phẩm đồng thời, hệ thống sẽ tuần tự hóa '
           'các thao tác và không để xảy ra tình trạng tồn kho bị âm.')

    add_heading_2(doc, '3.3. Tài liệu API và kiểm thử')
    add_bt(doc,
           'Toàn bộ 14 URL endpoint (tương ứng 23 thao tác HTTP — đếm theo '
           'cặp method và URL) được tài liệu hóa tự động qua Swagger UI tại '
           'đường dẫn /swagger/ nhờ thư viện drf-yasg. Bộ kiểm thử sử dụng '
           'Postman với file collection “ezware_postman_collection.json” gồm '
           'hơn 30 test case, có script tự lưu token sau bước login và tự '
           'điền vào các request tiếp theo. Kết quả kiểm thử được minh hoạ ở '
           'Phụ lục F (Hình 4–6). Các kịch bản đã được kiểm chứng bao gồm: '
           'đăng ký người dùng (với kiểm thử regression chống leo thang quyền), '
           'đăng nhập, tạo phiếu nhập, duyệt phiếu, tạo phiếu xuất với tồn '
           'kho thiếu (phải bị từ chối), hủy phiếu, và xem báo cáo tồn thấp.')


def build_design(doc):
    """H1 thứ 4 — THIẾT KẾ HỆ THỐNG."""
    add_heading_1(doc, '4. THIẾT KẾ HỆ THỐNG')

    add_heading_2(doc, '4.1. Sơ đồ lớp Model')
    add_bt(doc,
           'Mỗi bảng trong cơ sở dữ liệu được ánh xạ thành một lớp Model '
           'trong Django. Sơ đồ lớp được trình bày tại Hình 3.')
    add_figure_placeholder(doc, 'Hình 3. Sơ đồ lớp các Model của hệ thống EZWare.')
    add_bt(doc,
           'Các lớp Model: User, Product, Warehouse, InventoryReceipt, '
           'ReceiptDetail, Inventory. Mỗi lớp khai báo db_table tường minh '
           'để khớp tên bảng trong SQLite với tên trong tài liệu thiết kế. '
           'Các trường khóa ngoại sử dụng on_delete phù hợp với ngữ nghĩa: '
           'PROTECT cho dữ liệu master (warehouse, product) để ngăn chặn '
           'xóa bừa, CASCADE cho dữ liệu phụ thuộc (receipt details thuộc về '
           'một phiếu).')

    add_heading_2(doc, '4.2. Mẫu thiết kế đã áp dụng (Design patterns)')
    add_bt(doc,
           'Hệ thống áp dụng một số mẫu thiết kế phổ biến của Django REST '
           'Framework:')
    add_g1(doc, 'Class-based view: ListCreateAPIView, RetrieveUpdateDestroyAPIView, APIView — biến thể của Template Method pattern, giảm code lặp.')
    add_g1(doc, 'ModelSerializer: phân tách logic chuyển đổi giữa Model và JSON — biến thể Adapter pattern.')
    add_g1(doc, 'BasePermission: cho phép cắm-rút (plug-and-play) các chính sách phân quyền — Strategy pattern.')
    add_g1(doc, 'Manager (UserManager.create_user, create_superuser): mẫu Factory cho việc khởi tạo đối tượng User với mật khẩu được băm đúng cách.')

    add_heading_2(doc, '4.3. Áp dụng mô hình MVT (Model–View–Template)')
    add_bt(doc,
           'Django sử dụng biến thể MVT của mô hình MVC: Model phụ trách logic '
           'dữ liệu (ánh xạ ORM tới bảng SQLite), View phụ trách xử lý nghiệp '
           'vụ và trả về phản hồi (ở dạng JSON đối với REST API), Template ở '
           'dự án backend này không sử dụng vì đầu ra là JSON cho client. Hệ '
           'thống áp dụng nguyên tắc lập trình hướng đối tượng triệt để: toàn '
           'bộ View đều là class-based view kế thừa từ DRF, Serializer kế thừa '
           'từ ModelSerializer, Permission kế thừa từ BasePermission. Đồng '
           'thời, mỗi nghiệp vụ được tách thành một app Django riêng biệt '
           '(accounts, products, warehouses, inventory, reports) — đây là cách '
           'tổ chức giúp dễ bảo trì và mở rộng theo nguyên tắc Single '
           'Responsibility.')


def build_deployment(doc):
    """H1 thứ 5 — KIẾN TRÚC TRIỂN KHAI."""
    add_heading_1(doc, '5. KIẾN TRÚC TRIỂN KHAI')

    add_heading_2(doc, '5.1. Tổ chức thư mục source code')
    add_bt(doc,
           'Source code được tổ chức theo chuẩn project Django, mỗi nghiệp vụ '
           'là một app độc lập:')
    add_g1(doc, 'ezware/accounts — Quản lý user, xác thực, phân quyền.')
    add_g1(doc, 'ezware/products — Quản lý danh mục sản phẩm.')
    add_g1(doc, 'ezware/warehouses — Quản lý danh mục kho.')
    add_g1(doc, 'ezware/inventory — Phiếu nhập / xuất, tồn kho.')
    add_g1(doc, 'ezware/reports — Báo cáo tồn thấp.')
    add_g1(doc, 'ezware/core — Constants chung và lệnh import_excel.')
    add_g1(doc, 'ezware/settings.py — Cấu hình tập trung (DB, REST, Swagger).')

    add_heading_2(doc, '5.2. Triển khai chạy ở môi trường local và production')
    add_bt(doc,
           'Hệ thống được phát triển trên môi trường Windows với Python 3.10+ '
           'và đã được triển khai thực tế lên nền tảng Render.com (Python 3.12.3) '
           f'tại địa chỉ công khai {DEPLOY_URL}. Cụ thể: ở môi trường local '
           'sử dụng lệnh `python manage.py runserver` của Django; ở môi trường '
           'production trên Render sử dụng Gunicorn làm WSGI server với lệnh '
           '`gunicorn ezware.wsgi:application`, kèm WhiteNoise để phục vụ các '
           'file tĩnh (CSS / JS của Swagger UI). Database SQLite được sử dụng '
           'cho cả hai môi trường vì tính gọn nhẹ và phù hợp với scope đồ án '
           'môn học. Chi tiết các bước cài đặt local và deploy lên Render '
           'được trình bày tại Phụ lục E.')

    add_heading_2(doc, '5.3. Phụ thuộc thư viện')
    add_bt(doc,
           'Hệ thống dùng hai file requirements để tách biệt phụ thuộc '
           'production và development:')
    add_g1(doc, 'requirements.txt — Django 5.0.6, djangorestframework 3.15.1, drf-yasg 1.21.7, gunicorn 23.0.0, whitenoise 6.7.0, setuptools<81. Đủ để chạy server production.')
    add_g1(doc, 'requirements-dev.txt — kèm thêm pandas 2.2.2, numpy<2.0, openpyxl 3.1.2 để chạy lệnh import_excel nạp dữ liệu mẫu (chỉ cần ở local).')
    add_bt(doc,
           'Việc tách hai file giúp giảm dung lượng cài đặt trên server '
           'production (Render free tier chỉ có 512 MB RAM build) và phản '
           'ánh đúng nguyên tắc minimize attack surface.', first_line_indent=0)


def build_conclusion(doc):
    add_heading_1(doc, '6. KẾT LUẬN')
    add_bt(doc,
           'Sau quá trình phân tích, thiết kế và xây dựng, nhóm chúng tôi đã '
           'hoàn thiện một hệ thống backend API quản lý kho EZWare chạy ổn '
           'định, đáp ứng đầy đủ các yêu cầu cốt lõi: quản lý người dùng / '
           'sản phẩm / kho, lập và duyệt phiếu nhập–xuất, cập nhật tồn kho '
           'an toàn nhờ transaction nguyên tử, và báo cáo tồn kho thấp. Hệ '
           'thống được tổ chức theo mô hình MVT của Django, có phân quyền '
           'hai cấp và xác thực bằng Token; toàn bộ 14 endpoint được test '
           'bằng Postman với hơn 30 test case.')
    add_bt(doc,
           'Chức năng hài lòng nhất là nghiệp vụ duyệt phiếu xuất kho: hệ '
           'thống kiểm tra tồn trước, sau đó cập nhật trong transaction.atomic '
           'kết hợp select_for_update để chống race condition khi nhiều admin '
           'duyệt cùng lúc. Đây là chức năng có giá trị thực tế cao nhất đối '
           'với một hệ thống quản lý kho thực tế và minh hoạ rõ ràng năng lực '
           'lập trình đa luồng an toàn của Python với Django.')
    add_bt(doc,
           'Hạn chế: hệ thống mới chỉ chạy trên SQLite cho mục đích đồ án, '
           'chưa thử nghiệm trên PostgreSQL hay SQL Server với tải lớn; chưa '
           'có giao diện người dùng riêng (chỉ có Swagger UI cho thử nghiệm); '
           'chưa hỗ trợ các nghiệp vụ phức tạp hơn như chuyển kho trực tiếp '
           'giữa hai kho hay điều chỉnh tồn kho theo kiểm kê định kỳ.')


def build_references(doc):
    add_heading_1(doc, '7. TÀI LIỆU THAM KHẢO')
    today = datetime.now().strftime("%d/%m/%Y")
    add_tltk(doc, f'[1] Django Software Foundation, Django Documentation v5.0. Link: https://docs.djangoproject.com/en/5.0/ (Truy cập {today}).')
    add_tltk(doc, f'[2] Tom Christie và cộng sự, Django REST Framework Documentation. Link: https://www.django-rest-framework.org/ (Truy cập {today}).')
    add_tltk(doc, f'[3] OpenAPI Initiative, drf-yasg — Yet Another Swagger Generator. Link: https://drf-yasg.readthedocs.io/en/stable/ (Truy cập {today}).')
    add_tltk(doc, f'[4] SQLite Consortium, SQLite Documentation. Link: https://www.sqlite.org/docs.html (Truy cập {today}).')


# ============================================================
# Phụ lục
# ============================================================
def build_appendix_assignment(doc):
    add_centered(doc, 'PHỤ LỤC A. PHÂN CÔNG NHIỆM VỤ', bold=True, size=15)
    doc.add_paragraph()
    rows = [
        ('1', 'Nguyễn Văn A (25XXXXX1)',
         'Thiết kế cơ sở dữ liệu và viết các Model + migration; phát triển '
         'API quản lý sản phẩm và kho; viết Postman collection; chuẩn bị '
         'slide thuyết trình.'),
        ('2', 'Nguyễn Văn B (25XXXXX2)',
         'Phát triển API xác thực và phân quyền; hiện thực nghiệp vụ phiếu '
         'nhập / xuất với transaction.atomic; viết báo cáo Word; quay video '
         'thuyết trình.'),
    ]
    t = doc.add_table(rows=len(rows) + 1, cols=3)
    for i, h in enumerate(['STT', 'Thành viên', 'Nhiệm vụ']):
        t.rows[0].cells[i].text = h
    for i, r in enumerate(rows, start=1):
        for j, val in enumerate(r):
            t.rows[i].cells[j].text = val
    fix_table_font(t, size=12)


def build_appendix_summary(doc):
    add_centered(doc, 'PHỤ LỤC B. TÓM TẮT KẾT QUẢ', bold=True, size=15)
    add_centered(doc, '(Phụ lục bắt buộc theo Template)', italic=True, size=11)
    doc.add_paragraph()

    rows = [
        ('Chủ đề', 'Web API Application / Backend API Server (RESTful) — Chủ đề 2 theo đề bài.'),
        ('Giải pháp lập trình',
         'Mô hình MVT của Django (biến thể MVC); lập trình hướng đối tượng triệt để với class-based view, '
         'ModelSerializer, BasePermission, UserManager.'),
        ('Các loại user',
         'Quyền admin (user_role=admin): user admin01 / pass 123456 — toàn quyền. '
         'Quyền staff (user_role=staff): user nv01 / pass 123456 — xem danh mục, tạo phiếu.'),
        ('Số lượng bảng CSDL',
         '6 bảng: users, products, warehouses, inventory_receipts, receipt_details, inventory '
         '(thỏa yêu cầu 5–10 bảng của đồ án).'),
        ('Số lượng chức năng',
         '4 nhóm chính: Xác thực, Quản lý danh mục, Vận hành kho, Báo cáo. '
         'Chi tiết tại Mục 3.'),
        ('Chức năng hài lòng nhất',
         'Duyệt phiếu xuất kho — kiểm tra tồn trước, cập nhật trong transaction.atomic kết hợp '
         'select_for_update. Xem chi tiết tại Mục 3.2.'),
        ('Số lượng API',
         '14 URL endpoint với tổng 23 thao tác HTTP (mỗi URL có thể hỗ trợ '
         'nhiều method GET / POST / PUT / DELETE). Danh sách đầy đủ tại Phụ lục C.'),
        ('API hài lòng nhất',
         'PUT /api/receipts/<id>/status — duyệt / hủy phiếu, có atomic transaction. '
         'Xem chi tiết tại Mục 3.2.'),
        ('Tài liệu API',
         'Swagger UI tự sinh tại /swagger/ — hỗ trợ Token authentication và Security Definition.'),
        ('Kiểm thử',
         'Có. Postman collection “ezware_postman_collection.json” với hơn 30 test case, '
         'có script auto-save token sau login.'),
        ('Triển khai (deploy) công khai',
         f'Đã deploy thành công lên Render: {DEPLOY_URL}/swagger/ '
         '(Python 3.12.3 + Gunicorn + WhiteNoise). Bộ phận chấm có thể truy cập '
         'trực tiếp từ Internet để kiểm thử mà không cần cài đặt local.'),
        ('Công nghệ sử dụng',
         'Python 3.10+ (production 3.12.3), Django 5.0.6, Django REST Framework '
         '3.15.1, drf-yasg 1.21.7, gunicorn 23.0.0, whitenoise 6.7.0, SQLite, '
         'pandas + openpyxl (chỉ ở môi trường dev cho lệnh import_excel).'),
    ]

    t = doc.add_table(rows=len(rows) + 1, cols=3)
    for i, h in enumerate(['STT', 'Nội dung', 'Kết quả']):
        t.rows[0].cells[i].text = h
    for idx, (label, value) in enumerate(rows, start=1):
        t.rows[idx].cells[0].text = str(idx)
        t.rows[idx].cells[1].text = label
        t.rows[idx].cells[2].text = value
    fix_table_font(t, size=11)


def build_appendix_endpoints(doc):
    add_centered(doc, 'PHỤ LỤC C. DANH SÁCH ENDPOINT API', bold=True, size=15)
    doc.add_paragraph()
    add_bt(doc,
           'Bảng dưới đây liệt kê đầy đủ 14 URL endpoint của hệ thống (tổng '
           '23 thao tác HTTP đếm theo cặp method và URL), đáp ứng yêu cầu của '
           'Chủ đề 2 — Web API Application. Mỗi URL có thể hỗ trợ nhiều method, '
           'ví dụ /api/products có cả GET (list) và POST (create).',
           first_line_indent=0)

    endpoints = [
        # (Method, URL, Mô tả, Quyền)
        ('POST', '/api/auth/register', 'Đăng ký tài khoản mới, trả token', 'Public'),
        ('POST', '/api/auth/login', 'Đăng nhập, trả token', 'Public'),
        ('POST', '/api/auth/logout', 'Đăng xuất, xóa token', 'User đăng nhập'),
        ('GET / PUT', '/api/auth/me', 'Xem / cập nhật profile', 'User đăng nhập'),
        ('GET', '/api/products', 'Danh sách sản phẩm', 'User đăng nhập'),
        ('POST', '/api/products', 'Tạo sản phẩm', 'Admin'),
        ('GET / PUT / DELETE', '/api/products/<id>', 'Xem / sửa / xóa SP', 'Admin (PUT/DELETE)'),
        ('GET', '/api/warehouses', 'Danh sách kho', 'User đăng nhập'),
        ('POST', '/api/warehouses', 'Tạo kho', 'Admin'),
        ('GET / PUT / DELETE', '/api/warehouses/<id>', 'Xem / sửa / xóa kho', 'Admin (PUT/DELETE)'),
        ('GET', '/api/warehouses/<id>/inventory', 'Xem tồn kho theo kho', 'User đăng nhập'),
        ('GET / POST', '/api/receipts', 'Danh sách / tạo phiếu', 'User đăng nhập'),
        ('GET / DELETE', '/api/receipts/<id>', 'Xem / xóa phiếu (PENDING)', 'User đăng nhập / Admin'),
        ('POST', '/api/receipts/<id>/details', 'Thêm chi tiết phiếu', 'User đăng nhập'),
        ('PUT', '/api/receipts/<id>/status', 'Duyệt / hủy phiếu', 'Admin'),
        ('GET', '/api/reports/low-stock', 'Báo cáo tồn thấp', 'User đăng nhập'),
    ]
    t = doc.add_table(rows=len(endpoints) + 1, cols=4)
    for i, h in enumerate(['Method', 'URL', 'Mô tả', 'Quyền']):
        t.rows[0].cells[i].text = h
    for i, ep in enumerate(endpoints, start=1):
        for j, val in enumerate(ep):
            t.rows[i].cells[j].text = val
    fix_table_font(t, size=10)


def build_appendix_schema(doc):
    add_centered(doc, 'PHỤ LỤC D. CHI TIẾT SCHEMA CSDL', bold=True, size=15)
    doc.add_paragraph()
    add_bt(doc,
           'Phụ lục này mô tả chi tiết từng cột của 6 bảng trong cơ sở dữ liệu '
           'EZWare. Tất cả tên cột được đặt theo quy ước snake_case.',
           first_line_indent=0)

    tables_schema = [
        ('users', [
            ('user_id', 'INTEGER PK AutoField', 'Khóa chính tự sinh'),
            ('username', 'VARCHAR(150) UNIQUE', 'Tên đăng nhập, không trùng'),
            ('password', 'VARCHAR', 'Mật khẩu băm PBKDF2 (Django default)'),
            ('full_name', 'VARCHAR(200)', 'Họ tên đầy đủ'),
            ('email', 'VARCHAR(200)', 'Email (optional)'),
            ('phone', 'VARCHAR(20)', 'Số điện thoại'),
            ('user_role', 'VARCHAR(10)', 'Vai trò: admin / staff'),
            ('is_active', 'BOOLEAN', 'Tài khoản còn hoạt động'),
            ('is_staff', 'BOOLEAN', 'Quyền vào Django admin (auto-sync từ user_role qua save())'),
            ('last_login', 'DATETIME', 'Lần đăng nhập gần nhất (từ AbstractBaseUser)'),
        ]),
        ('products', [
            ('product_id', 'INTEGER PK AutoField', 'Khóa chính tự sinh'),
            ('product_code', 'VARCHAR(50) UNIQUE', 'Mã sản phẩm'),
            ('product_name', 'VARCHAR(200)', 'Tên sản phẩm'),
            ('product_type', 'VARCHAR(100)', 'Loại / nhóm sản phẩm'),
            ('product_description', 'TEXT', 'Mô tả'),
            ('is_active', 'BOOLEAN', 'Còn kinh doanh hay đã tạm ngưng'),
        ]),
        ('warehouses', [
            ('warehouse_id', 'INTEGER PK AutoField', 'Khóa chính'),
            ('warehouse_name', 'VARCHAR(200) UNIQUE', 'Tên kho'),
            ('warehouse_location', 'VARCHAR(500)', 'Địa chỉ kho'),
            ('is_active', 'BOOLEAN', 'Đang hoạt động'),
        ]),
        ('inventory_receipts', [
            ('receipt_id', 'INTEGER PK AutoField', 'Khóa chính'),
            ('warehouse_id', 'INTEGER FK', 'Kho áp dụng — FK warehouses, on_delete=PROTECT'),
            ('receipt_type', 'VARCHAR(10)', 'IMPORT / EXPORT'),
            ('receipt_status', 'VARCHAR(15)', 'PENDING / APPROVED / CANCELLED (mặc định PENDING)'),
            ('created_at', 'DATETIME', 'Thời gian tạo phiếu (auto_now_add)'),
        ]),
        ('receipt_details', [
            ('detail_id', 'INTEGER PK AutoField', 'Khóa chính'),
            ('receipt_id', 'INTEGER FK', 'Phiếu cha — FK inventory_receipts, on_delete=CASCADE'),
            ('product_id', 'INTEGER FK', 'Sản phẩm — FK products, on_delete=PROTECT'),
            ('quantity', 'PositiveInteger', 'Số lượng. Validate min=1 ở tầng serializer'),
        ]),
        ('inventory', [
            ('id', 'INTEGER PK AutoField', 'Khóa chính tự sinh (Django auto)'),
            ('warehouse_id', 'INTEGER FK', 'Kho — FK warehouses, on_delete=CASCADE'),
            ('product_id', 'INTEGER FK', 'Sản phẩm — FK products, on_delete=CASCADE'),
            ('quantity', 'INTEGER (default=0)', 'Số lượng tồn thực tế'),
            ('—', 'UNIQUE constraint', 'UNIQUE(warehouse_id, product_id) — mỗi SP 1 dòng tồn / kho'),
        ]),
    ]
    for table_name, cols in tables_schema:
        p = doc.add_paragraph()
        run = p.add_run(f'Bảng {table_name}:')
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        t = doc.add_table(rows=len(cols) + 1, cols=3)
        for i, h in enumerate(['Cột', 'Kiểu / ràng buộc', 'Ý nghĩa']):
            t.rows[0].cells[i].text = h
        for i, (c, ty, mean) in enumerate(cols, start=1):
            t.rows[i].cells[0].text = c
            t.rows[i].cells[1].text = ty
            t.rows[i].cells[2].text = mean
        fix_table_font(t, size=10)
        doc.add_paragraph()


def build_appendix_install(doc):
    add_centered(doc, 'PHỤ LỤC E. HƯỚNG DẪN CÀI ĐẶT VÀ CHẠY', bold=True, size=15)
    doc.add_paragraph()
    add_bt(doc,
           'Phụ lục này gồm hai phần: (E.1–E.5) hướng dẫn chạy ở máy local '
           'và (E.6) hướng dẫn deploy lên Render. Yêu cầu hệ thống: '
           'Python 3.10+ và pip.', first_line_indent=0)

    add_heading_2(doc, 'E.1. Cài đặt thư viện')
    add_bt(doc,
           'Cài đặt bằng requirements-dev.txt (kèm pandas/openpyxl để nạp Excel):',
           first_line_indent=0)
    p = doc.add_paragraph()
    run = p.add_run('pip install -r requirements-dev.txt')
    run.font.name = 'Consolas'; run.font.size = Pt(11)

    add_heading_2(doc, 'E.2. Khởi tạo database')
    p = doc.add_paragraph()
    run = p.add_run('python manage.py makemigrations accounts products warehouses inventory\npython manage.py migrate')
    run.font.name = 'Consolas'; run.font.size = Pt(11)

    add_heading_2(doc, 'E.3. Nạp dữ liệu mẫu (chọn 1 trong 2)')
    add_bt(doc, 'Cách 1 — nạp từ Excel mẫu kèm theo:', first_line_indent=0)
    p = doc.add_paragraph()
    run = p.add_run('python manage.py import_excel EZWare_DuLieuMau.xlsx --reset')
    run.font.name = 'Consolas'; run.font.size = Pt(11)
    add_bt(doc, 'Cách 2 — tạo admin trống rồi tự thêm data qua API:', first_line_indent=0)
    p = doc.add_paragraph()
    run = p.add_run(
        "python manage.py shell -c \"from ezware.accounts.models import User; "
        "User.objects.create_user(username='admin01', password='123456', "
        "user_role='admin', full_name='Quản trị viên')\""
    )
    run.font.name = 'Consolas'; run.font.size = Pt(10)

    add_heading_2(doc, 'E.4. Chạy server local')
    p = doc.add_paragraph()
    run = p.add_run('python manage.py runserver')
    run.font.name = 'Consolas'; run.font.size = Pt(11)
    add_bt(doc, 'Sau khi server chạy, truy cập:', first_line_indent=0)
    add_g1(doc, 'http://127.0.0.1:8000/swagger/ — tài liệu API tự sinh.')
    add_g1(doc, 'http://127.0.0.1:8000/admin/ — Django admin.')
    add_g1(doc, 'Postman: import file ezware_postman_collection.json để chạy bộ test case.')

    add_heading_2(doc, 'E.5. Tài khoản mẫu sau import')
    add_g1(doc, 'admin01 / 123456 — quyền admin (toàn quyền).')
    add_g1(doc, 'nv01, nv02 / 123456 — quyền staff (chỉ đọc danh mục, tạo phiếu).')

    add_heading_2(doc, 'E.6. Deploy lên Render (đã thực hiện)')
    add_bt(doc,
           f'Hệ thống đã được deploy công khai tại {DEPLOY_URL}/swagger/. '
           'Các bước thực hiện:', first_line_indent=0)
    add_g1(doc, 'Push source code lên GitHub.')
    add_g1(doc, 'Tạo Web Service trên https://render.com, kết nối với repository.')
    add_g1(doc, 'Build Command: pip install -r requirements.txt && python manage.py collectstatic --no-input && python manage.py migrate')
    add_g1(doc, 'Start Command: gunicorn ezware.wsgi:application')
    add_g1(doc, 'Thiết lập biến môi trường DJANGO_SECRET_KEY (random), DJANGO_DEBUG=False, DJANGO_ALLOWED_HOSTS=<tên-app>.onrender.com')
    add_bt(doc,
           'Lưu ý kỹ thuật: file requirements.txt phục vụ production cố tình '
           'không kèm pandas/numpy/openpyxl để tránh tràn 512 MB RAM build '
           'của Render free tier; pin setuptools<81 do drf-yasg 1.21.7 còn '
           'import pkg_resources mà setuptools 81+ đã loại bỏ module này; '
           'sử dụng WhiteNoise để phục vụ static files (CSS/JS Swagger UI) '
           'trên môi trường production mà không cần Nginx.')


def build_appendix_figures(doc):
    add_centered(doc, 'PHỤ LỤC F. HÌNH KIỂM THỬ', bold=True, size=15)
    doc.add_paragraph()
    add_bt(doc,
           'Phụ lục này thể hiện kết quả kiểm thử các API quan trọng của hệ '
           'thống thông qua công cụ Postman và Swagger UI. Mỗi hình bao gồm: '
           'URL endpoint thực tế trên Render, request body, response status '
           f'code và response body. Tất cả các request đều gọi đến {DEPLOY_URL} '
           '— hệ thống đang chạy thực trên Internet, người chấm có thể tự '
           'kiểm chứng bằng cách mở Postman / browser.',
           first_line_indent=0)

    add_figure_placeholder(doc, 'Hình 4. Kiểm thử API đăng nhập trên Postman — request POST /api/auth/login và token trả về.')
    add_figure_placeholder(doc, 'Hình 5. Kiểm thử luồng tạo phiếu nhập → thêm chi tiết → duyệt → xem tồn kho cập nhật.')
    add_figure_placeholder(doc, 'Hình 6. Giao diện Swagger UI tại /swagger/ hiển thị 14 URL endpoint với 23 thao tác HTTP của hệ thống.')


def build_appendix_code(doc):
    add_centered(doc, 'PHỤ LỤC G. CODE TRÍCH (TUỲ CHỌN)', bold=True, size=15)
    doc.add_paragraph()
    add_bt(doc,
           'Phụ lục này trích phần code quan trọng nhất — nghiệp vụ duyệt '
           'phiếu trong transaction nguyên tử. Toàn bộ source code được nộp '
           'kèm trong thư mục Nhom XX_Source.',
           first_line_indent=0)
    code = (
        '# ezware/inventory/views.py — trích đoạn approve receipt\n'
        'with transaction.atomic():\n'
        '    for ct in chi_tiet_list:\n'
        '        ton, _ = Inventory.objects.select_for_update().get_or_create(\n'
        '            warehouse_id=phieu.warehouse_id,\n'
        '            product_id=ct.product_id,\n'
        '            defaults={"quantity": 0},\n'
        '        )\n'
        '        if phieu.receipt_type == RECEIPT_TYPE_IMPORT:\n'
        '            ton.quantity += ct.quantity\n'
        '        else:\n'
        '            ton.quantity -= ct.quantity\n'
        '        ton.save(update_fields=["quantity"])\n'
        '\n'
        '    phieu.receipt_status = RECEIPT_STATUS_APPROVED\n'
        '    phieu.save(update_fields=["receipt_status"])\n'
    )
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)


# ============================================================
# Main build
# ============================================================
def main():
    doc = Document()

    # Section 1: trang bìa, không footer/page-number
    s1 = doc.sections[0]
    s1.top_margin = Cm(2.0)
    s1.bottom_margin = Cm(2.0)
    s1.left_margin = Cm(2.5)
    s1.right_margin = Cm(2.0)

    create_custom_styles(doc)

    # ---- Trang bìa ----
    build_cover(doc)

    # ---- Section 2: nội dung chính (có footer) ----
    s2 = add_section_page_break(doc)
    s2.footer.is_linked_to_previous = False
    setup_footer(s2, MEMBERS_SHORT)

    build_intro(doc)
    build_database(doc)
    build_features(doc)
    build_design(doc)
    build_deployment(doc)
    build_conclusion(doc)
    build_references(doc)

    # ---- Phụ lục ----
    for builder in [
        build_appendix_assignment,
        build_appendix_summary,
        build_appendix_endpoints,
        build_appendix_schema,
        build_appendix_install,
        build_appendix_figures,
        build_appendix_code,
    ]:
        add_page_break(doc)
        builder(doc)

    doc.save(OUTPUT_NAME)
    print(f'Đã sinh file: {OUTPUT_NAME}')
    print('Các bước tiếp theo sinh viên cần làm:')
    print('  1) Thay placeholder Nhóm XX, Nguyễn Văn A/B, MSSV trong file')
    print('  2) Sửa footer (chuột phải Edit Footer) — đổi MEMBERS_SHORT')
    print('  3) Chèn 3 sơ đồ: ER (Hình 1), chức năng (Hình 2), lớp (Hình 3)')
    print('  4) Chèn 3 screenshot Postman/Swagger (Hình 4-6) vào Phụ lục F')
    print('  5) Save as PDF để nộp kèm bản docx')


if __name__ == '__main__':
    main()
